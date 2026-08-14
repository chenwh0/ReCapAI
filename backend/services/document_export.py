from docx import Document
from fpdf import FPDF
from pathlib import Path
from ..models.schemas import RecapaiFormat

def export_recap(recap: RecapaiFormat, export_type: str):
    export_types = {".txt": save_as_txt, ".docx": save_as_docx, ".pdf": save_as_pdf}
    if export_type in export_types:
        return export_types[export_type](recap)
    else: 
        raise ValueError(f"Unsupported export type: {export_type}")


def save_as_txt(recap: RecapaiFormat, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headers & newlines as simple text document(.txt)
        Parameters:
        recap -- RecapaiFormat(dict), Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save txt to (default=./output_recap/)
        file_name -- str, name of txt (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = Path(file_path) / f"{file_name}.txt"
    file_address.parent.mkdir(parents=True, exist_ok=True)

    with open(file_address, "w", encoding="utf-8") as file:
        for key, value in recap.model_dump().items():
            file.write(f"{key.upper()}\n")
            if key != "SENTIMENT":
                file.write(f"{value}\n\n")
            else:
                file.write(f"{value.sentiment_label}\n")
                file.write(f"{value.explanation}\n")
    print(f"Saving {file_name}.txt at {file_path}")
    return file_address

def save_as_docx(recap: RecapaiFormat, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headings & newlines as Microsoft Word docx
        Parameters:
        recap -- RecapaiFormat(dict), Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save docx to (default=./output_recap/)
        file_name -- str, name of docx (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = Path(file_path) / f"{file_name}.docx"
    file_address.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("SUMMARY", level=1)
    doc.add_paragraph(recap.summary)
    doc.add_paragraph()

    doc.add_heading("KEY POINTS", level=1)
    for point in recap.key_points:
        doc.add_paragraph(point)
    doc.add_paragraph()

    doc.add_heading("TASKS", level=1)
    for task in recap.task:
        doc.add_paragraph(task)
    doc.add_paragraph()

    doc.add_heading("SENTIMENT", level=1)
    doc.add_paragraph(recap.sentiment.sentiment_label)
    doc.add_paragraph(recap.sentiment.explanation)

    doc.save(file_address)
    return file_address

def save_as_pdf(recap: RecapaiFormat, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headings & newlines as pdf
        Parameters:
        recap -- RecapaiFormat(dict), Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save pdf to (default=./output_recap/)
        file_name -- str, name of pdf (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = Path(file_path) / f"{file_name}.pdf"
    file_address.parent.mkdir(parents=True, exist_ok=True)
    pdf = FPDF()
    pdf.add_page()
    for key, value in recap.model_dump().items():
        # Write key names as large cornflower blue headers in pdfs
        pdf.set_font("Helvetica", style="B", size=32)
        pdf.set_text_color(100, 149, 237) # cornflower blue
        pdf.write(16, f"{key}\n")

        # Write values (responses) as paragraph text in pdf
        pdf.set_font("Helvetica",size=12)
        pdf.set_text_color(0, 0, 0) # black
        pdf.multi_cell(0, 8, value)
        pdf.write(5, "\n\n")
    pdf.output(file_address)
    return file_address